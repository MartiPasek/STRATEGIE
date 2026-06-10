# -*- coding: utf-8 -*-
"""Generátor personálních dokumentů ze živých dat STRATEGIE (Marti 10.6.2026 — „v ERP na klik").
Vstup = data dict (z PG engagement/wage + Helios TabCisZam/účet), výstup = docx (bytes).
Typy: smlouva | vymer | popis | dpp. Univerzální: firma z COMPANY, text generický,
[DOPLNIT: …] červeně. Brand: logo (static) + patička + Verdana.
"""
from __future__ import annotations
import io
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOPL = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x80, 0x80, 0x80)
OPEN_Q = "„"   # „
CLOSE_Q = "“"  # "

# logo: apps/api/static/brand/eurosoft_logo.png (přidáno do repa)
_HERE = os.path.dirname(os.path.abspath(__file__))
LOGO_CANDIDATES = [
    os.path.join(_HERE, "..", "..", "..", "apps", "api", "static", "brand", "eurosoft_logo.png"),
    "/opt/strategie/apps/api/static/brand/eurosoft_logo.png",
]

def _logo_path():
    for p in LOGO_CANDIDATES:
        if os.path.exists(p):
            return os.path.abspath(p)
    return None

# Univerzální konfigurace firem (prodejné: nová firma = nový záznam)
COMPANY = {
    "ES": dict(nazev="EUROSOFT - System s.r.o.", sidlo="Nepomucká 259, 326 00 Plzeň",
               ico="26411741", dic="CZ26411741",
               oR="Krajský soud v Plzni, oddíl C, vložka 18532",
               jednatel="Martin Pašek, jednatel",
               tel_absence="773 738 585 nebo 777 180 511"),
    "EC": dict(nazev="EUROSOFT - Control s.r.o.", sidlo="Nepomucká 259, 326 00 Plzeň",
               ico="27960862", dic="CZ27960862",
               oR="Krajský soud v Plzni, oddíl C, vložka 18759",
               jednatel="Martin Pašek, jednatel a Branislav Mózer, jednatel",
               tel_absence="773 738 585 nebo 777 180 511"),
}
MESICE = ["", "ledna", "února", "března", "dubna", "května", "června", "července",
          "srpna", "září", "října", "listopadu", "prosince"]
TODAY = None  # nastaví endpoint (cz datum)


def cz_date(iso):
    if not iso:
        return None
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", str(iso))
    if not m:
        return str(iso)
    y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
    return "%d. %s %d" % (d, MESICE[mo], y)


def money(n):
    try:
        return "{:,}".format(int(round(float(n)))).replace(",", " ") + " Kč"
    except Exception:
        return str(n)


def _sv(run):
    run.font.name = "Verdana"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), "Verdana")


def _brand(doc, co):
    st = doc.styles["Normal"]
    st.font.name = "Verdana"
    st.font.size = Pt(10.5)
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp = _logo_path()
    if lp:
        try:
            hp.add_run().add_picture(lp, width=Inches(1.5))
        except Exception:
            r = hp.add_run(co["nazev"]); r.bold = True; _sv(r)
    else:
        r = hp.add_run(co["nazev"]); r.bold = True; _sv(r)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("%s  |  %s  |  IČ: %s  |  www.eurosoft.com" % (co["nazev"], co["sidlo"], co["ico"]))
    fr.font.size = Pt(7.5); _sv(fr); fr.font.color.rgb = GREY


def _H(doc, t, sz=15):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = True; r.font.size = Pt(sz); _sv(r)


def _sub(doc):
    p = doc.add_paragraph()
    r = p.add_run("NÁVRH vygenerovaný ze systému STRATEGIE — k doplnění a kontrole")
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY; _sv(r)


def _P(doc, t="", bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    for part in re.split(r"(\[DOPLNIT[^\]]*\])", t):
        if not part:
            continue
        r = p.add_run(part); _sv(r)
        if part.startswith("[DOPLNIT"):
            r.bold = True; r.font.color.rgb = DOPL
        else:
            r.bold = bold
    return p


def _v(x, ph):
    """vrať text hodnoty, nebo [DOPLNIT: ph]"""
    if x is None or str(x).strip() == "" or "[DOPLNIT" in str(x):
        return "[DOPLNIT: %s]" % ph
    return str(x)


def _osoba_blok(doc, d, role="zaměstnanec"):
    jm = (str(d.get("jmeno") or "").strip())
    _P(doc, "%s narozen/a %s" % (_v(jm, "jméno a příjmení"), _v(cz_date(d.get("narozeni")), "datum narození")))
    _P(doc, "Trvalé bydliště: " + _v(d.get("bydliste"), "trvalé bydliště"))
    _P(doc, "– dále jen %s –" % role)


def _podpis(doc, co, jmeno, role="zaměstnanec"):
    _P(doc, "")
    _P(doc, "V Plzni dne " + (TODAY or "[DOPLNIT: datum]"))
    _P(doc, "")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "______________________________"
    t.rows[0].cells[1].text = "______________________________"
    a = t.rows[1].cells[0].paragraphs[0]; a.add_run("%s, zaměstnavatel" % co["nazev"])
    b = t.rows[1].cells[1].paragraphs[0]; b.add_run("%s, %s" % (_v(jmeno, "jméno"), role))
    for cell in (t.rows[1].cells[0], t.rows[1].cells[1]):
        for r in cell.paragraphs[0].runs:
            _sv(r)


def build_smlouva(doc, d, co):
    _brand(doc, co)
    _H(doc, "Pracovní smlouva"); _sub(doc)
    _P(doc, "Mezi " + co["nazev"]); _P(doc, "Sídlo: " + co["sidlo"]); _P(doc, "IČ: " + co["ico"])
    _P(doc, "– dále jen zaměstnavatel –"); _P(doc, "a")
    _osoba_blok(doc, d)
    _P(doc, "se uzavírá následující pracovní smlouva:")
    _P(doc, "§ 1 Doba a obsah pracovního poměru", bold=True)
    _P(doc, "1. Zaměstnanec nastupuje od %s jako %s." % (_v(cz_date(d.get("od")), "den nástupu"), _v(d.get("pozice"), "druh práce / pozice")))
    if d.get("do"):
        _P(doc, "2. Pracovní smlouva se sjednává na dobu určitou do %s." % cz_date(d.get("do")))
    else:
        _P(doc, "2. Pracovní smlouva se sjednává na dobu neurčitou.")
    _P(doc, "§ 2 Místo výkonu práce a pracovní doba", bold=True)
    _P(doc, "1. Místem výkonu práce je Plzeň.")
    uv = d.get("uvazek")
    _P(doc, "2. Pracovní doba činí %s hodin týdně." % (("%.0f" % float(uv)) if uv else "[DOPLNIT: úvazek]"))
    _P(doc, "3. Zaměstnanec zahájí práci nejpozději v 9:00, konec 14:00–18:00, obvykle pondělí–pátek; po dohodě i o víkendu.")
    _P(doc, "4. Práce přesčas max. 150 hodin za kalendářní rok; k přesčasu nad 30 min denně se přihlédne při mimořádných odměnách.")
    _P(doc, "§ 3 Mzda", bold=True)
    _P(doc, "Zaměstnanci náleží mzda dle mzdového výměru. Mzda je sjednána již s přihlédnutím k případné práci přesčas.")
    _P(doc, "§ 4 Další příjmy", bold=True)
    _P(doc, "Prémie, zvláštní platby a odměny závisí na rozhodnutí vedení firmy; není na ně právní nárok.")
    _P(doc, "§ 5 Pracovní schopnost", bold=True)
    _P(doc, "Zaměstnanec je povinen oznámit nepřítomnost předem, nejpozději v 9:00, na číslo " + co["tel_absence"] + ".")
    _P(doc, "§ 6–11", bold=True)
    _P(doc, "Ukončení dle zákoníku práce; změny písemně; mlčenlivost (i po skončení); konkurenční doložka dle přílohy č. 1; hmotná odpovědnost; prohlášení o seznámení s předpisy.")
    _P(doc, "§ 12 Dovolená", bold=True)
    _P(doc, "Čerpání dovolené se řídí ustanoveními zákoníku práce České republiky.")
    _P(doc, "§ 13 Obecná ustanovení", bold=True)
    _P(doc, "Vnitřní předpisy: Etický kodex a firemní kultura; konto pracovní doby; GDPR I. a II. část. Zaměstnanec podpisem potvrzuje seznámení a akceptaci.")
    _podpis(doc, co, d.get("jmeno"))


def build_vymer(doc, d, co):
    _brand(doc, co)
    today = TODAY or "[DOPLNIT: datum]"
    _H(doc, "Mzdový výměr", 14)
    _P(doc, "ze dne " + today)
    _P(doc, "mezi " + co["nazev"]); _P(doc, "Sídlo: " + co["sidlo"] + "  IČ: " + co["ico"])
    _P(doc, "- dále jen zaměstnavatel –"); _P(doc, "a")
    _osoba_blok(doc, d)
    _P(doc, "")
    _P(doc, "1. S platností od %s obdrží zaměstnanec za svou činnost:" % today, bold=True)
    _P(doc, "    • hrubou měsíční mzdu ve výši %s," % _v(money(d.get("zaklad")) if d.get("zaklad") else None, "základní mzda"))
    _P(doc, "    • osobní ohodnocení ve výši od 0 Kč do %s, dle rozhodnutí zaměstnavatele a s ohledem na body 2. a 3." % _v(money(d.get("os_ohod")) if d.get("os_ohod") else None, "horní hranice os. ohodnocení"))
    for label, val in (d.get("extra") or []):
        _P(doc, "    • %s: %s" % (label, money(val)))
    uv = d.get("uvazek")
    if uv and float(uv) < 40:
        _P(doc, "Pracovní doba činí %.0f hodin týdně; mzda odpovídá délce kratší pracovní doby." % float(uv))
    _P(doc, "Částky jsou splatné do 17. dne následujícího měsíce na konto č. %s." % _v(d.get("ucet"), "číslo účtu"))
    _P(doc, "")
    _P(doc, "2. Osobní ohodnocení bude vyplaceno v plné výši při řádném plnění úkolů a dodržování pracovní smlouvy a vnitřních předpisů (Etický kodex, Firemní kultura).")
    _P(doc, "3. Při nedostatku zakázek nebo platební neschopnosti odběratelů si zaměstnavatel vyhrazuje právo na snížení nebo nevyplacení osobního ohodnocení.")
    _podpis(doc, co, d.get("jmeno"))


def build_popis(doc, d, co):
    _brand(doc, co)
    p = doc.add_paragraph(); r = p.add_run("§ 15  Popis pracovního místa"); r.bold = True; r.font.size = Pt(14); _sv(r)
    _sub(doc)
    _P(doc, "Pozice: " + _v(d.get("pozice"), "pozice"), bold=True)
    _P(doc, "")
    _P(doc, "Produkt:", bold=True)
    for b in ("kvalitní a důsledné provedení svěřené práce ve sjednaném termínu",
              "dodržení správných pracovních postupů a předpisů BOZP",
              "spolupráce v týmu a aktivní řešení nejasností"):
        _P(doc, "    • " + b)
    _P(doc, "")
    _P(doc, "Odborné a osobnostní předpoklady:", bold=True)
    for b in ("vůle k optimální spolupráci na pracovišti (teamwork)",
              "práce zaměřená na úspěšný výsledek a spokojenost zaměstnavatele",
              "časová flexibilita, vlastní iniciativa, spolehlivost a odpovědnost",
              "hospodárné myšlení, slušné chování, diskrétnost"):
        _P(doc, "    • " + b)
    kat = (d.get("kategorie") or "").strip()
    if kat:
        _P(doc, "")
        _P(doc, "Kategorie: " + kat, bold=True)
    _P(doc, "")
    _P(doc, "Nejvyšší prioritou je korektní jednání v týmu a kvalitní provedení svěřené práce ve smluveném termínu.")
    _podpis(doc, co, d.get("jmeno"))


def build_dpp(doc, d, co):
    _brand(doc, co)
    _H(doc, "Dohoda o provedení práce"); _sub(doc)
    _P(doc, "mezi " + co["nazev"]); _P(doc, "Sídlo: " + co["sidlo"] + "  IČ: " + co["ico"])
    _P(doc, "– dále jen zaměstnavatel –"); _P(doc, "a")
    _osoba_blok(doc, d)
    _P(doc, "se uzavírá následující dohoda o provedení práce:")
    _P(doc, "1. Zaměstnanec provede tyto práce: %s." % _v(d.get("pozice"), "vymezení práce"))
    _P(doc, "2. Práce v období od %s; rozsah a doba dle domluvy s jednatelem." % _v(cz_date(d.get("od")), "datum"))
    _P(doc, "3. Odměna dle odpracovaných hodin výkonovou mzdou ve výši [DOPLNIT: sazba Kč/h].")
    _P(doc, "4. Odměna splatná po ukončení zakázek na konto č. %s." % _v(d.get("ucet"), "číslo účtu"))
    _P(doc, "5. Rozsah max. 300 hodin za kalendářní rok. 6. Práci provede osobně.")
    _P(doc, "Vnitřní předpisy (Etický kodex, konto pracovní doby, GDPR) zaměstnanec podpisem akceptuje.")
    _podpis(doc, co, d.get("jmeno"))


BUILDERS = {"smlouva": build_smlouva, "vymer": build_vymer, "popis": build_popis, "dpp": build_dpp}
TITLES = {"smlouva": "pracovni_smlouva", "vymer": "mzdovy_vymer", "popis": "popis_pracovniho_mista", "dpp": "DPP"}


def generate(data: dict, typ: str):
    """data: jmeno, narozeni, bydliste, firma(EC/ES), pozice, od, do, uvazek,
    zaklad, os_ohod, extra[(label,val)], ucet, kategorie. Vrací (filename, bytes)."""
    co = COMPANY.get((data.get("firma") or "ES").upper(), COMPANY["ES"])
    builder = BUILDERS.get(typ)
    if builder is None:
        raise ValueError("neznámý typ dokumentu: %s" % typ)
    doc = Document()
    builder(doc, data, co)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    nm = re.sub(r"[^A-Za-z0-9ÁČĎÉĚÍŇÓŘŠŤÚŮÝŽáčďéěíňóřšťúůýž_-]", "", (data.get("jmeno") or "zamestnanec").replace(" ", "_"))
    fname = "%s_%s.docx" % (nm, TITLES.get(typ, typ))
    return fname, buf.getvalue()
