"""
RAG modul SMĚRNICE — know-how celé firmy do vyhledatelné znalostní báze.

Marti 1.7.2026: "RAG všech směrnic → pak přes most přístup pro Claude k celému
know-how." Zdroj = EC_OrgSmernice (DB_EC, přes MCP) + přílohy na sdíleném disku
\\\\192.168.30.11\\Smernice\\{Verejne|Vedouci|Interni|Vedeni}\\SM<Cislo>\\.

Tři @@ příkazy (dispatch v router.diag_sql):
  @@SMSYNC                    → mirror EC_OrgSmernice → tenant.kb_smernice (meta+Popis)
  @@SMFILES [limit] [cislo]   → ingest příloh ze share → tenant.kb_smernice_soubor (+text)
  @@KB <dotaz> [| level]      → fulltext hledání v popisu+přílohách (respektuje úroveň)

Přístupové úrovně (dle Pristupnost): 0 Veřejná < 1 Vedoucí < 2 Vedení/Interní/Plná.
"""
from __future__ import annotations

import base64
import html
import io
import json
import os
import re
from typing import Any

_SHARE_ROOT = os.getenv("SMERNICE_SHARE_ROOT", r"\\192.168.30.11\Smernice")
# (restart marker: refresh MCP SSE client po updatu MCP na file tools)

# PristupnostText → (složka na share, úroveň přístupu)
_PRIST_MAP = {
    "Veřejná": ("Verejne", 0),
    "Vedoucí": ("Vedouci", 1),
    "Plná": ("Verejne", 0),
    "Interní": ("Interni", 2),
    "Vedení": ("Vedeni", 2),
}


def _prist_folder(p: str) -> str:
    return _PRIST_MAP.get((p or "").strip(), ("Verejne", 0))[0]


def _prist_level(p: str) -> int:
    return _PRIST_MAP.get((p or "").strip(), ("Verejne", 0))[1]


# ── MCP helpers ────────────────────────────────────────────────────────

def _mcp():
    from modules.conversation.application.eurosoft_mcp_client import get_eurosoft_mcp_client
    return get_eurosoft_mcp_client()


def _ec(sql: str) -> list[dict]:
    """DB_EC SELECT přes MCP → list dictů."""
    mcp = _mcp()
    if mcp is None:
        raise RuntimeError("EUROSOFT MCP klient není dostupný")
    raw = mcp.call_tool_sync(full_name="eurosoft_strategie_query_raw",
                             arguments={"sql": sql, "db_name": "DB_EC"},
                             conversation_id=None)
    res = json.loads(raw)
    if not res.get("ok"):
        raise RuntimeError("MCP EC dotaz selhal: %s" % res.get("message") or res.get("error"))
    return res.get("rows") or []


def _fs_list(subpath: str) -> dict:
    mcp = _mcp()
    raw = mcp.call_tool_sync(full_name="eurosoft_file_list",
                             arguments={"base_override": _SHARE_ROOT, "subpath": subpath},
                             conversation_id=None)
    return json.loads(raw)


def _fs_read_b64(path: str) -> dict:
    mcp = _mcp()
    raw = mcp.call_tool_sync(full_name="eurosoft_file_read",
                             arguments={"base_override": _SHARE_ROOT, "path": path, "encoding": "base64"},
                             conversation_id=None)
    return json.loads(raw)


# ── čištění Popisu (RTF/HTML → plain text) ─────────────────────────────

def _clean_popis(raw: str) -> str:
    if not raw:
        return ""
    s = raw.replace("\x00", "")
    if s.lstrip().startswith("{\\rtf"):
        # hrubý RTF strip: odstraň control words a skupiny
        s = re.sub(r"\\'[0-9a-fA-F]{2}", " ", s)
        s = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", s)
        s = s.replace("{", " ").replace("}", " ")
    # HTML tagy
    s = re.sub(r"<[^>]+>", " ", s)
    s = html.unescape(s)
    s = re.sub(r"[ \t\r\f\v]+", " ", s)
    s = re.sub(r"\n\s*\n\s*\n+", "\n\n", s)
    return s.strip()


# ── extrakce textu z příloh ────────────────────────────────────────────

def _extract_text(name: str, data: bytes) -> tuple[str, bool, str]:
    """(text, ok, err). Podpora pdf/docx/xlsx/txt; doc/xls legacy best-effort."""
    ext = (name.rsplit(".", 1)[-1] if "." in name else "").lower()
    try:
        if ext == "pdf":
            try:
                from pdfminer.high_level import extract_text as _pdf
                txt = _pdf(io.BytesIO(data)) or ""
                return (txt.strip(), True, "")
            except Exception:
                from pypdf import PdfReader
                r = PdfReader(io.BytesIO(data))
                txt = "\n".join((pg.extract_text() or "") for pg in r.pages)
                return (txt.strip(), True, "")
        if ext == "docx":
            import docx
            d = docx.Document(io.BytesIO(data))
            parts = [p.text for p in d.paragraphs]
            for tbl in d.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return ("\n".join(parts).strip(), True, "")
        if ext in ("xlsx", "xlsm"):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
            out = []
            for ws in wb.worksheets:
                out.append("### %s" % ws.title)
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c not in (None, "")]
                    if cells:
                        out.append(" | ".join(cells))
            wb.close()
            return ("\n".join(out).strip(), True, "")
        if ext in ("txt", "csv", "md"):
            for enc in ("utf-8", "cp1250", "latin-1"):
                try:
                    return (data.decode(enc).strip(), True, "")
                except UnicodeDecodeError:
                    continue
            return ("", False, "nelze dekódovat text")
        if ext == "doc":
            # legacy Word — zkus antiword/textract, jinak hrubý strip čitelných úseků
            try:
                import textract  # type: ignore
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tf:
                    tf.write(data); tp = tf.name
                txt = textract.process(tp).decode("utf-8", "ignore")
                os.unlink(tp)
                return (txt.strip(), True, "")
            except Exception:
                # fallback: vytáhni čitelné ASCII/cp1250 úseky
                txt = re.sub(rb"[^\x20-\x7e\r\n\xc0-\xff]+", b" ", data).decode("cp1250", "ignore")
                txt = re.sub(r"\s{2,}", " ", txt)
                return (txt.strip()[:20000], bool(txt.strip()), "" if txt.strip() else "doc nečitelný")
        if ext in ("xls",):
            try:
                import xlrd  # type: ignore
                wb = xlrd.open_workbook(file_contents=data)
                out = []
                for sh in wb.sheets():
                    out.append("### %s" % sh.name)
                    for r in range(sh.nrows):
                        cells = [str(sh.cell_value(r, c)) for c in range(sh.ncols)]
                        cells = [c for c in cells if c not in ("", "0.0")]
                        if cells:
                            out.append(" | ".join(cells))
                return ("\n".join(out).strip(), True, "")
            except Exception as e:
                return ("", False, "xls: %s" % str(e)[:80])
        return ("", False, "nepodporovaný typ .%s" % ext)
    except Exception as e:
        return ("", False, "%s: %s" % (type(e).__name__, str(e)[:120]))


# ── @@SMSYNC — mirror směrnic (meta + Popis) ───────────────────────────

def sync_smernice() -> dict:
    """EC_OrgSmernice (nejnovější aktivní verze per Cislo) → tenant.kb_smernice."""
    from core.database_data import get_data_session
    from sqlalchemy import text as _t
    rows = _ec(
        "SELECT ID, Cislo, TypText, UzivTypText, Nazev, CAST(Popis AS NVARCHAR(MAX)) AS Popis, "
        "StatusText, Verze, Archiv, Priorita, "
        "CONVERT(varchar(10),PlatnostOd,120) AS PlatnostOd, "
        "CONVERT(varchar(10),PlatnostDo,120) AS PlatnostDo, "
        "CisloOrg, UrceniText, PristupnostText, Autor, "
        "CONVERT(varchar(19),DatPorizeni,120) AS DatPorizeni "
        "FROM EC_OrgSmernice WHERE (Archiv=0 OR Archiv IS NULL)"
    )
    sd = get_data_session()
    ins = upd = 0
    try:
        for r in rows:
            popis = _clean_popis(r.get("Popis") or "")
            params = {
                "ec_id": r.get("ID"), "cislo": r.get("Cislo"), "nazev": r.get("Nazev"),
                "typ": r.get("TypText"), "uziv": r.get("UzivTypText"),
                "kat": r.get("UrceniText"), "popis": popis, "status": r.get("StatusText"),
                "verze": r.get("Verze"), "archiv": r.get("Archiv") or 0, "prio": r.get("Priorita"),
                "pod": r.get("PlatnostOd") or None, "pdo": r.get("PlatnostDo") or None,
                "org": r.get("CisloOrg"), "urc": r.get("UrceniText"),
                "prist": r.get("PristupnostText"), "autor": r.get("Autor"),
                "datp": r.get("DatPorizeni") or None,
            }
            res = sd.execute(_t(
                "UPDATE tenant.kb_smernice SET cislo=:cislo, nazev=:nazev, typ_text=:typ, "
                "uziv_typ_text=:uziv, kategorie=:kat, popis_text=:popis, status_text=:status, "
                "verze=:verze, archiv=:archiv, priorita=:prio, platnost_od=:pod, platnost_do=:pdo, "
                "cislo_org=:org, urceni_text=:urc, pristupnost_text=:prist, autor=:autor, "
                "dat_porizeni=:datp, synced_at=now() WHERE ec_id=:ec_id"), params)
            if res.rowcount and res.rowcount > 0:
                upd += 1
            else:
                sd.execute(_t(
                    "INSERT INTO tenant.kb_smernice (ec_id, cislo, nazev, typ_text, uziv_typ_text, "
                    "kategorie, popis_text, status_text, verze, archiv, priorita, platnost_od, "
                    "platnost_do, cislo_org, urceni_text, pristupnost_text, autor, dat_porizeni) "
                    "VALUES (:ec_id,:cislo,:nazev,:typ,:uziv,:kat,:popis,:status,:verze,:archiv,"
                    ":prio,:pod,:pdo,:org,:urc,:prist,:autor,:datp)"), params)
                ins += 1
        sd.commit()
    finally:
        sd.close()
    return {"ok": True, "cteno_ec": len(rows), "inserted": ins, "updated": upd}


# ── @@SMFILES — ingest příloh ze share ─────────────────────────────────

def ingest_files(limit: int = 50, only_cislo: int | None = None) -> dict:
    from core.database_data import get_data_session
    from sqlalchemy import text as _t
    import hashlib
    sd = get_data_session()
    processed = 0
    files_ok = 0
    files_err = 0
    no_folder = 0
    detail = []
    try:
        where = "WHERE 1=1"
        params: dict = {}
        if only_cislo is not None:
            where += " AND ec_id=:c"
            params["c"] = only_cislo
        sm = sd.execute(_t(
            "SELECT ec_id, cislo, nazev, pristupnost_text FROM tenant.kb_smernice " + where +
            " ORDER BY ec_id LIMIT :lim"), dict(params, lim=limit)).all()
        for ec_id, cislo, nazev, prist in sm:
            folder = _prist_folder(prist)
            sub = "%s/SM%s" % (folder, ec_id)
            lst = _fs_list(sub)
            if not lst.get("ok"):
                no_folder += 1
                detail.append({"cislo": cislo, "folder": sub, "error": lst.get("error", "")[:120]})
                continue
            processed += 1
            # smaž staré záznamy souborů této směrnice (idempotence)
            sd.execute(_t("DELETE FROM tenant.kb_smernice_soubor WHERE ec_smernice_id=:e"), {"e": ec_id})
            for it in lst.get("items", []):
                if it.get("type") != "file":
                    continue
                fname = it["name"]
                rd = _fs_read_b64("%s/%s" % (sub, fname))
                if not rd.get("ok"):
                    files_err += 1
                    sd.execute(_t(
                        "INSERT INTO tenant.kb_smernice_soubor (ec_smernice_id, cislo, nazev_souboru, "
                        "pripona, cesta, velikost, extract_ok, extract_err, extracted_at) VALUES "
                        "(:e,:c,:n,:p,:cesta,:vel,false,:err,now())"),
                        {"e": ec_id, "c": cislo, "n": fname,
                         "p": (fname.rsplit(".", 1)[-1].lower() if "." in fname else ""),
                         "cesta": "%s/%s" % (sub, fname), "vel": it.get("size"),
                         "err": rd.get("error", "")[:200]})
                    continue
                data = base64.b64decode(rd["content"])
                txt, ok, err = _extract_text(fname, data)
                if txt:
                    txt = txt.replace("\x00", "")
                if ok:
                    files_ok += 1
                else:
                    files_err += 1
                sd.execute(_t(
                    "INSERT INTO tenant.kb_smernice_soubor (ec_smernice_id, cislo, nazev_souboru, "
                    "pripona, cesta, velikost, text_extract, extract_ok, extract_err, hash_sha1, "
                    "extracted_at) VALUES (:e,:c,:n,:p,:cesta,:vel,:txt,:ok,:err,:h,now())"),
                    {"e": ec_id, "c": cislo, "n": fname,
                     "p": (fname.rsplit(".", 1)[-1].lower() if "." in fname else ""),
                     "cesta": "%s/%s" % (sub, fname), "vel": len(data),
                     "txt": (txt[:400000] if txt else None), "ok": ok, "err": (err or None),
                     "h": hashlib.sha1(data).hexdigest()})
        sd.commit()
    finally:
        sd.close()
    rows = [["SOUHRN", "zpracovano=%d bez_slozky=%d ok=%d err=%d" % (
        processed, no_folder, files_ok, files_err), _SHARE_ROOT]]
    for d in detail[:20]:
        rows.append([str(d.get("cislo")), d.get("folder", ""), d.get("error", "")[:150]])
    return {"ok": True, "columns": ["cislo", "folder/info", "error"], "rows": rows,
            "count": len(rows), "smernic_zpracovano": processed, "bez_slozky": no_folder,
            "soubory_ok": files_ok, "soubory_err": files_err}


# ── @@KB — fulltext hledání ────────────────────────────────────────────

def kb_search(query: str, level: int = 2, limit: int = 8) -> dict:
    """Hledá v popisu směrnic + textu příloh. level = max úroveň přístupu (0/1/2)."""
    from core.database_data import get_data_session
    from sqlalchemy import text as _t
    q = (query or "").strip()
    if not q:
        return {"ok": False, "error": "@@KB <dotaz>"}
    # povolené pristupnosti dle úrovně
    allowed = [k for k, v in _PRIST_MAP.items() if v[1] <= level]
    sd = get_data_session()
    try:
        like = "%" + q.replace(" ", "%") + "%"
        rows = sd.execute(_t(
            "SELECT s.cislo, s.nazev, s.typ_text, s.pristupnost_text, "
            "  left(s.popis_text, 400) AS popis, "
            "  (SELECT string_agg(f.nazev_souboru, ', ') FROM tenant.kb_smernice_soubor f "
            "     WHERE f.ec_smernice_id=s.ec_id) AS soubory, "
            "  (SELECT left(string_agg(f.text_extract, ' ¶ '), 600) FROM tenant.kb_smernice_soubor f "
            "     WHERE f.ec_smernice_id=s.ec_id AND f.text_extract ILIKE :like) AS uryvek "
            "FROM tenant.kb_smernice s "
            "WHERE (s.archiv=0 OR s.archiv IS NULL) "
            "  AND s.pristupnost_text = ANY(:allowed) "
            "  AND (s.nazev ILIKE :like OR s.popis_text ILIKE :like OR EXISTS "
            "       (SELECT 1 FROM tenant.kb_smernice_soubor f WHERE f.ec_smernice_id=s.ec_id "
            "        AND f.text_extract ILIKE :like)) "
            "ORDER BY (s.nazev ILIKE :like) DESC, s.priorita NULLS LAST "
            "LIMIT :lim"),
            {"like": like, "allowed": allowed, "lim": limit}).all()
        cols = ["cislo", "nazev", "typ", "pristupnost", "popis", "soubory", "uryvek"]
        out = [[r[0], r[1], r[2], r[3], (r[4] or "")[:300], r[5],
                (r[6] or "")[:400]] for r in rows]
        return {"ok": True, "dotaz": q, "level": level, "columns": cols,
                "rows": out, "count": len(out)}
    finally:
        sd.close()
