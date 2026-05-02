"""
Phase 27e (2.5.2026 ráno): DOCX reader pro Marti-AI.

Po Klárka workflow Phase 27 epoch (1.5.2026): Marti chce dotáhnout
business stack -- Word DOCX read. Po dnešku Marti-AI umí Excel + PDF +
Image + DOCX read; plus sandbox pro generování čehokoliv.

Marti-AI's design choices (2.5.2026 ráno, Phase 27e konzultace):
  A - Output structured (paragraphs + tables + metadata, analog Excel/PDF)
  A - Headings v paragraphs s typed metadata (type='heading', level=N)
  A - Vše dostupné metadata + aproximace word_count
  A - Legacy .doc -> error s návodem (uložte jako .docx)
  + insider vstup: prazdne paragrafy = "esteticke mezery", default skip
    (parametr include_empty_paragraphs: bool = False)

Implementační poznámky:
  - python-docx jen pro .docx (modern Word XML format)
  - Pro .doc (legacy binární Word 97-2003) error s helpful hláškou
  - core_properties NEvrací word_count -- aproximace ze split() per paragraph
  - page_count nelze ze python-docx (Word page rendering je layout-time)
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from core.logging import get_logger

logger = get_logger("rag.docx")

# Marti-AI's volby z Phase 27e konzultace
ALLOWED_DOCX_TYPES = {"docx"}  # python-docx jen modern XML format
LEGACY_DOC_HINT = (
    "Document je legacy .doc (Word 97-2003). python-docx neumi tento format. "
    "Uložte v Wordu jako .docx (Soubor → Uložit jako → Word Document .docx) "
    "a nahrajte znovu."
)


def _resolve_document_path(document_id: int) -> tuple[str, str, int | None, str]:
    """Vrátí (storage_path, original_filename, tenant_id, ext) pro DOCX document."""
    from core.database import get_session
    from modules.core.infrastructure.models_data import Document

    session = get_session()
    try:
        doc = session.query(Document).filter_by(id=document_id).first()
        if not doc:
            raise ValueError(f"Document id={document_id} nenalezen.")
        if not doc.storage_path:
            raise ValueError(f"Document id={document_id} nema storage_path.")

        ext = (doc.file_type or "").lower().lstrip(".")
        if ext == "doc":
            raise ValueError(LEGACY_DOC_HINT)
        if ext not in ALLOWED_DOCX_TYPES:
            raise ValueError(
                f"Document id={document_id} neni DOCX (file_type='{doc.file_type}'). "
                "Pro Excel pouzij read_excel_structured, pro PDF read_pdf_structured."
            )

        return (
            doc.storage_path,
            doc.original_filename or doc.name or f"document_{document_id}",
            doc.tenant_id,
            ext,
        )
    finally:
        session.close()


def _check_tenant_access(document_tenant_id: int | None, caller_tenant_id: int | None, is_parent: bool) -> bool:
    """Standardní tenant gate s parent bypass."""
    if is_parent:
        return True
    if document_tenant_id is None:
        return True  # inbox / global
    if caller_tenant_id is None:
        return False
    return document_tenant_id == caller_tenant_id


def _extract_metadata(core_properties: Any, doc: Any) -> dict:
    """
    Extract metadata z python-docx core_properties + spočítej word_count.

    Marti-AI's volba A: vše dostupné. core_properties.author / title /
    subject / keywords / category / created / last_modified / revision.
    Plus word_count aproximace ze paragraph.text.split().
    """
    md: dict = {}

    try:
        md["author"] = core_properties.author or None
    except Exception:
        md["author"] = None
    try:
        md["title"] = core_properties.title or None
    except Exception:
        md["title"] = None
    try:
        md["subject"] = core_properties.subject or None
    except Exception:
        md["subject"] = None
    try:
        md["keywords"] = core_properties.keywords or None
    except Exception:
        md["keywords"] = None
    try:
        md["category"] = core_properties.category or None
    except Exception:
        md["category"] = None
    try:
        md["created"] = core_properties.created.isoformat() if core_properties.created else None
    except Exception:
        md["created"] = None
    try:
        md["last_modified"] = (
            core_properties.modified.isoformat() if core_properties.modified else None
        )
    except Exception:
        md["last_modified"] = None
    try:
        md["revision"] = core_properties.revision
    except Exception:
        md["revision"] = None

    # Word count aproximace -- python-docx nedrzi page_count, ale paragraph text je
    # k dispozici. Ne dokonale (ne pocita words v tabulkach), ale dobry signal pro
    # Marti-AI's rozhodnuti "cist cely doc nebo jen prvnich N paragraphs".
    word_count = 0
    try:
        for p in doc.paragraphs:
            if p.text:
                word_count += len(p.text.split())
        # Plus words v tabulkach
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        word_count += len(cell.text.split())
    except Exception as e:
        logger.warning(f"DOCX word_count failed: {e}")
    md["word_count"] = word_count

    return md


def _classify_paragraph(paragraph: Any) -> tuple[str, int | None]:
    """
    Klasifikuje paragraph na typ + heading level.

    Word style names: "Heading 1" / "Heading 2" / ... / "Title" / "Normal" /
    "List Paragraph" / "Quote" / atd.

    Returns: (type, level)
      ("heading", N) pro Heading 1-9 + Title (level 0)
      ("paragraph", None) pro běžný text
      ("empty", None) pro prazdne paragraphs (Marti-AI's insight)
    """
    text = (paragraph.text or "").strip()
    style_name = ""
    try:
        style_name = paragraph.style.name or ""
    except Exception:
        pass

    if not text:
        return ("empty", None)

    style_lower = style_name.lower()
    if style_lower.startswith("heading "):
        # "Heading 1" -> level 1
        try:
            level = int(style_lower.split()[-1])
            if 1 <= level <= 9:
                return ("heading", level)
        except (ValueError, IndexError):
            pass
        return ("heading", None)
    if style_lower == "title":
        return ("heading", 0)  # Title = top-level (above Heading 1)

    return ("paragraph", None)


def _extract_table(table: Any) -> list[list[str]]:
    """Word table -> list[list[str]] (rows of cells)."""
    out = []
    try:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append((cell.text or "").strip())
            out.append(row_data)
    except Exception as e:
        logger.warning(f"DOCX table extract failed: {e}")
    return out


def read_docx_structured(
    document_id: int,
    *,
    include_empty_paragraphs: bool = False,
    caller_tenant_id: int | None = None,
    is_parent: bool = False,
) -> dict:
    """
    Vrátí structured Word DOCX (paragraphs + tables + metadata).

    Args:
        document_id: ID dokumentu z RAG documents (file_type='docx')
        include_empty_paragraphs: Marti-AI's insight (Phase 27e konzultace) --
            Word docs mají hodně prázdných paragraphs jako "estetické mezery".
            Default False = tiše skipnout (čistá data). True = include
            s type='empty' (pro debug nebo když user chce kompletní strukturu).
        caller_tenant_id, is_parent: standardní tenant gate

    Output:
        {
          "document_id": 42,
          "filename": "smlouva.docx",
          "metadata": {
            "author": "Marti Pasek",
            "title": "Najemni smlouva",
            "last_modified": "2026-04-15T10:30:00",
            "word_count": 1234,
            ...
          },
          "paragraphs": [
            {"index": 0, "type": "heading", "level": 1, "text": "Najemni smlouva"},
            {"index": 1, "type": "paragraph", "text": "Pronajimatel..."},
            ...
          ],
          "tables": [
            [["Kolonka", "Hodnota"], ["Cena", "15000 Kc/mesic"]]
          ],
          "n_paragraphs": 42,
          "n_tables": 3,
          "warnings": []
        }
    """
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise RuntimeError(f"python-docx neni dostupna: {e}")

    storage_path, filename, doc_tenant_id, _ext = _resolve_document_path(document_id)

    if not _check_tenant_access(doc_tenant_id, caller_tenant_id, is_parent):
        raise PermissionError(
            f"Document id={document_id} patri jinemu tenantu (doc.tenant_id={doc_tenant_id})."
        )

    p = Path(storage_path)
    if not p.is_file():
        raise ValueError(f"Soubor neexistuje na disku: {storage_path}")

    warnings: list[str] = []

    try:
        doc = DocxDocument(str(p))
    except Exception as e:
        msg = str(e).lower()
        if "package" in msg or "zip" in msg or "not a word" in msg:
            raise ValueError(
                "Document neni platny .docx (nebo je poskozeny). "
                "Zkontroluj že je to opravdu Word DOCX format. "
                f"python-docx error: {e}"
            )
        raise ValueError(f"DOCX parse failed: {type(e).__name__}: {e}")

    # Metadata
    try:
        metadata = _extract_metadata(doc.core_properties, doc)
    except Exception as e:
        warnings.append(f"metadata extrakce selhala: {type(e).__name__}: {e}")
        metadata = {}

    # Paragraphs
    paragraphs_out: list[dict] = []
    skipped_empty = 0
    for idx, para in enumerate(doc.paragraphs):
        ptype, level = _classify_paragraph(para)
        if ptype == "empty":
            if not include_empty_paragraphs:
                skipped_empty += 1
                continue
            paragraphs_out.append({
                "index": idx,
                "type": "empty",
                "text": "",
            })
            continue
        entry = {
            "index": idx,
            "type": ptype,
            "text": para.text.strip(),
        }
        if ptype == "heading" and level is not None:
            entry["level"] = level
        paragraphs_out.append(entry)

    if skipped_empty > 0:
        warnings.append(
            f"Skipped {skipped_empty} prazdnych paragraphs (esteticke mezery). "
            "Pro inkluzi nastav include_empty_paragraphs=True."
        )

    # Tables
    tables_out: list[list[list[str]]] = []
    for table in doc.tables:
        try:
            tables_out.append(_extract_table(table))
        except Exception as e:
            warnings.append(f"table extrakce selhala: {type(e).__name__}: {e}")

    return {
        "document_id": document_id,
        "filename": filename,
        "metadata": metadata,
        "paragraphs": paragraphs_out,
        "tables": tables_out,
        "n_paragraphs": len(paragraphs_out),
        "n_tables": len(tables_out),
        "warnings": warnings,
    }
